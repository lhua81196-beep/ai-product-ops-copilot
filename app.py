# -*- coding: utf-8 -*-
"""
AI Product Ops Copilot
基于 Streamlit + DeepSeek API 的产品运营助手。
"""

import os
from typing import Literal

import streamlit as st
from dotenv import load_dotenv

import pandas as pd
import io
import PyPDF2
from datetime import datetime

from api import chat_stream, chat_json
from rag_engine import VectorStore, split_text
from prompts import (
    competitive_analysis_system,
    competitive_analysis_user,
    competitive_matrix_system,
    competitive_matrix_user,
    knowledge_base_system,
    knowledge_base_user,
    campaign_planning_system,
    campaign_planning_user,
    marketing_copy_system,
    marketing_copy_user,
    weekly_report_system,
    weekly_report_user,
)

# ---------------------------------------------------------------------------
# 初始化
# ---------------------------------------------------------------------------

load_dotenv()

st.set_page_config(
    page_title="AI Product Ops Copilot",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# 自定义样式
# ---------------------------------------------------------------------------

st.markdown("""
<style>
    /* 全局字体与间距 */
    .block-container { padding-top: 2rem; padding-bottom: 2rem; }
    h1, h2, h3 { margin-top: 0; }

    /* Sidebar 品牌区 */
    .sidebar-brand {
        text-align: center;
        padding: 1.5rem 0.5rem 0.5rem;
    }
    .sidebar-brand h2 {
        font-size: 1.25rem;
        margin-bottom: 0.25rem;
    }
    .sidebar-brand p {
        font-size: 0.8rem;
        color: #9aa0a6;
        margin: 0;
    }

    /* 模块卡片式标题 */
    .module-header {
        margin-bottom: 1.5rem;
    }
    .module-header h1 {
        font-size: 1.75rem;
        font-weight: 600;
    }
    .module-header p {
        color: #5f6368;
        font-size: 0.95rem;
    }

    /* 结果容器 */
    .result-container {
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        padding: 1.5rem;
        margin-top: 1.5rem;
        background: #fafafa;
    }

    /* 状态提示 */
    .stAlert { border-radius: 6px; }

    @media (prefers-color-scheme: dark) {
        .module-header p { color: #9aa0a6; }
        .result-container {
            border-color: #333;
            background: #1e1e1e;
        }
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Sidebar 导航
# ---------------------------------------------------------------------------

def render_sidebar() -> Literal["competitive_analysis", "campaign_planning",
                                 "marketing_copy", "weekly_report",
                                 "knowledge_base"]:
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-brand">
            <h2>🚀 Product Ops</h2>
            <p>AI 驱动 · 高效运营</p>
        </div>
        """, unsafe_allow_html=True)

        st.divider()

        nav_options = {
            "competitive_analysis": "📊 竞品分析",
            "knowledge_base":      "📚 竞品知识库",
            "campaign_planning":   "🎯 活动策划",
            "marketing_copy":      "✍️ 营销文案",
            "weekly_report":       "📋 周报生成",
        }

        # 用 radio 做导航
        selected = st.radio(
            "功能选择",
            options=list(nav_options.keys()),
            format_func=lambda k: nav_options[k],
            label_visibility="collapsed",
        )

        st.divider()

        # 使用指南
        with st.expander("📌 使用指南", expanded=False):
            st.markdown("""
**📊 竞品分析**
1. 输入 3~5 个竞品名称（每行一个）
2. 点击 **开始分析**
3. 等待矩阵和 SWOT 结构化输出
4. 下载 **Word 报告**

**🎯 活动策划**
输入产品 + 目标 → 获取完整活动方案

**✍️ 营销文案**
输入产品 + 平台 → 生成适配文案

**📋 周报生成**
描述工作内容 → 自动生成周报

**📚 竞品知识库**
上传 PDF → 提问 → 基于资料回答
""")

        st.divider()

        # 历史分析记录
        st.session_state.setdefault("history", [])
        if st.session_state["history"]:
            with st.expander("📚 历史分析记录", expanded=False):
                for record in st.session_state["history"][:10]:
                    st.markdown(f"**{record['time']}**")
                    st.markdown(f"产品: {record['products'].replace(chr(10), ', ')[:50]}")
                    st.caption(f"摘要: {record['summary'][:120]}...")
                    st.divider()

        # API 密钥配置
        with st.expander("⚙️ API 配置", expanded=False):
            current_key = os.getenv("DEEPSEEK_API_KEY", "")
            masked = (current_key[:8] + "..." + current_key[-4:]
                      if len(current_key) > 12 else "")
            if masked:
                st.success(f"已配置: {masked}")
            else:
                api_key = st.text_input(
                    "DeepSeek API Key",
                    type="password",
                    placeholder="sk-...",
                    help="输入后保存在 Session 中，不会写入文件。",
                )
                if api_key:
                    os.environ["DEEPSEEK_API_KEY"] = api_key
                    st.rerun()

        # 页脚
        st.markdown(
            "<p style='text-align:center;color:#9aa0a6;font-size:0.75rem;'>"
            "v1.0 · Powered by DeepSeek</p>",
            unsafe_allow_html=True,
        )

    return selected


# ---------------------------------------------------------------------------
# 各模块渲染函数
# ---------------------------------------------------------------------------

def clean_markdown(text: str) -> str:
    """清洗 LLM 输出的 Markdown，移除残留标记。"""
    import re
    text = text.replace("```markdown", "")
    text = text.replace("```", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _set_east_asian(run) -> None:
    """为 run 设置东亚字体（确保中文正常显示）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def _add_run(p, text: str, bold: bool = False) -> None:
    """向 paragraph 添加一个带 Microsoft YaHei 字体的 run。"""
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.bold = bold
    _set_east_asian(run)


def _add_bold_runs(p, text: str) -> None:
    """向已有 paragraph 添加 runs，拆解 **bold** 标记。"""
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            _add_run(p, part[2:-2], bold=True)
        elif part:
            _add_run(p, part)
    if not any(r.text for r in p.runs):
        p._element.getparent().remove(p._element)


def _add_bold_paragraph(doc, text: str) -> None:
    """新建段落并添加 runs，拆解 **bold** 标记。"""
    p = doc.add_paragraph()
    _add_bold_runs(p, text)


def _add_markdown_content(doc, text: str) -> None:
    """将 Markdown 文本逐行解析为 Word 结构化段落/标题/列表。"""
    import re
    for raw_line in text.split('\n'):
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_runs(p, line[2:])
        elif re.match(r'^\d+[\.\)、]?\s', line):
            text = re.sub(r'^\d+[\.\)、]?\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_bold_runs(p, text)
        elif line.startswith('|'):
            continue
        else:
            _add_bold_paragraph(doc, line)


def generate_word_report(products: str, df: pd.DataFrame,
                         report_text: str) -> io.BytesIO:
    """生成 Word 格式竞品分析报告，返回 BytesIO 对象。"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement

    doc = Document()

    # ----- 全局字体：微软雅黑 -----
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for lvl in range(1, 4):
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.color.rgb = None  # 继承默认色
        rPr2 = hs.element.get_or_add_rPr()
        rFonts2 = rPr2.find(qn('w:rFonts'))
        if rFonts2 is None:
            rFonts2 = OxmlElement('w:rFonts')
            rPr2.insert(0, rFonts2)
        rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ----- 标题 -----
    title = doc.add_heading('竞品分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ----- 1. 分析对象 -----
    doc.add_heading('1. 分析对象', level=1)
    for name in (x.strip() for x in products.split('\n') if x.strip()):
        p = doc.add_paragraph(style='List Bullet')
        _add_run(p, name)

    # ----- 2. 竞品对比矩阵 -----
    doc.add_heading('2. 竞品对比矩阵', level=1)

    # 固定列顺序 + 清洗空值
    FIXED_COLUMNS = [
        "product", "positioning", "target_user",
        "core_advantage", "core_disadvantage", "business_model",
        "confidence",
        "confidence_reason",
    ]
    available = [c for c in FIXED_COLUMNS if c in df.columns]
    df_clean = df.reindex(columns=available).fillna("").astype(str)

    table = doc.add_table(rows=len(df_clean) + 1, cols=len(df_clean.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 均匀列宽
    total_width = Cm(16)
    col_w = total_width / len(df_clean.columns)
    for col_idx in range(len(df_clean.columns)):
        for cell in table.columns[col_idx].cells:
            cell.width = col_w

    header_map = {
        'product': '产品名称', 'positioning': '产品定位',
        'target_user': '目标用户', 'core_advantage': '核心优势',
        'core_disadvantage': '核心劣势', 'business_model': '商业模式',
        'confidence': '可信度',
    }

    # 表头行
    for j, col in enumerate(df_clean.columns):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, header_map.get(col, col), bold=True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F81BD"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # 白色字体
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
            _set_east_asian(run)

    # 数据行
    for i, row in enumerate(df_clean.itertuples(index=False)):
        for j, value in enumerate(row):
            # 安全转换 + 防超长截断
            safe = str(value).strip() if value is not None else ''
            if len(safe) > 2000:
                safe = safe[:2000] + '…[已截断]'
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            _add_run(cell.paragraphs[0], safe)
            # 设置 Arial 字体
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Arial'
                _set_east_asian(run)

    # ----- 3. SWOT 分析报告（Markdown → Word）-----
    doc.add_heading('3. SWOT 分析报告', level=1)
    _add_markdown_content(doc, report_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def render_competitive_analysis():
    st.markdown('<div class="module-header">'
                '<h1>📊 竞品分析</h1>'
                '<p>输入多个产品名称，生成竞品对比矩阵（表格）和完整 SWOT 分析报告。</p>'
                '</div>', unsafe_allow_html=True)

    with st.expander("📌 关于本工具", expanded=True):
        st.markdown("""
**AI 驱动的竞品分析工具**

帮助产品/运营团队快速完成竞品调研、结构化分析、SWOT 输出和报告生成。

**核心价值**
将 **2 小时**分析压缩到 **10 分钟**

**适用人群**
产品经理 · 产品运营 · 增长 / 市场人员

**分析维度**
产品定位 · 用户群体 · 核心优势 · 商业模式
""")

    products = st.text_area(
        "输入产品名称（每行一个）",
        placeholder="例：\nChatGPT\nClaude\nGemini",
        height=120,
    )

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        temperature = st.slider("创意度", 0.0, 1.5, 0.8, 0.1, key="atemp")

    with col2:
        max_tokens = st.selectbox("最大长度", [2048, 4096, 8192],
                                  index=1, key="atokens")

    if st.button("🚀 开始分析", type="primary", use_container_width=True):
        if not products.strip():
            st.warning("👉 请先输入竞品名称（建议3-5个）")
            return

        status = st.status("分析中...", expanded=True)

        # --- Phase 1: 竞品矩阵 ---
        status.update(label="📊 正在生成竞品矩阵...", state="running")
        _EMPTY_ROW = {
            "product": "", "positioning": "", "target_user": "",
            "core_advantage": "", "core_disadvantage": "", "business_model": "",
            "confidence": 0.0,
            "confidence_reason": "",
        }
        matrix_data = None
        try:
            matrix_data = chat_json(
                competitive_matrix_system(),
                competitive_matrix_user(products.strip()),
                temperature=temperature,
                max_tokens=2048,
            )
        except Exception:
            st.warning("⚠️ 模型返回异常，已使用空矩阵占位，请重试。")

        if not matrix_data:
            matrix_data = [_EMPTY_ROW]

        df = pd.DataFrame(matrix_data)
        FIXED_COLUMNS = [
            "product", "positioning", "target_user",
            "core_advantage", "core_disadvantage", "business_model",
            "confidence",
            "confidence_reason",
        ]
        available = [c for c in FIXED_COLUMNS if c in df.columns]
        df = df.reindex(columns=available).fillna("").astype(str)

        # --- Phase 2: SWOT 分析 ---
        status.update(label="🧠 正在构建SWOT分析...", state="running")
        collected = ""
        placeholder = st.empty()
        for chunk in chat_stream(
            competitive_analysis_system(),
            competitive_analysis_user(products.strip()),
            temperature=temperature,
            max_tokens=max_tokens,
        ):
            collected += chunk
            placeholder.markdown(clean_markdown(collected))

        status.update(label="✅ 分析完成", state="complete")

        # --- 保存到 Session ---
        st.session_state["report_products"] = products.strip()
        st.session_state["report_df_json"] = df.to_dict(orient="records")
        st.session_state["report_text"] = collected
        st.session_state["report_ready"] = True
        st.session_state.setdefault("analysis_history", [])
        st.session_state["analysis_history"].insert(0, {
            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "products": products.strip(),
            "dataframe": st.session_state["report_df_json"],
            "report_text": collected,
        })
        # 同步保存到 sidebar 历史
        st.session_state.setdefault("history", [])
        st.session_state["history"].insert(0, {
            "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "products": products.strip(),
            "summary": collected[:200],
        })
        st.rerun()


    if st.session_state.get("report_ready"):
        st.divider()

        with st.container():
            st.subheader("📊 结构化数据")
            df = pd.DataFrame.from_dict(st.session_state["report_df_json"])

            # 可信度映射
            if "confidence" in df.columns:
                df["confidence"] = df["confidence"].astype(float)
                if "confidence_reason" in df.columns:
                    df["confidence_reason"] = df["confidence_reason"].fillna("")

            st.dataframe(
                df,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "product":       "产品名称",
                    "positioning":   "产品定位",
                    "target_user":   "目标用户",
                    "core_advantage": "核心优势",
                    "core_disadvantage": "核心劣势",
                    "business_model": "商业模式",
                    "confidence": st.column_config.ProgressColumn(
                        "可信度",
                        help="基于公开信息和产品成熟度的可信度评分（0~1）",
                        min_value=0,
                        max_value=1,
                        format="%.0f%%",
                    ),
                    "confidence_reason": "可信度依据",
                },
            )

        st.divider()

        with st.container():
            st.subheader("🧠 AI 分析")
            st.markdown(clean_markdown(st.session_state["report_text"]))

        st.divider()

        with st.container():
            st.subheader("📄 导出区")
            df = pd.DataFrame.from_dict(st.session_state["report_df_json"])
            csv = df.to_csv(index=False).encode("utf-8-sig")
            buf = generate_word_report(
                st.session_state["report_products"],
                df,
                st.session_state["report_text"],
            )
            col_csv, col_word = st.columns(2)
            with col_csv:
                st.download_button(
                    label="📥 下载 CSV",
                    data=csv,
                    file_name="competitive_matrix.csv",
                    mime="text/csv",
                )
            with col_word:
                st.download_button(
                    label="📄 导出 Word 报告",
                    data=buf,
                    file_name="competitive_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


    # --- 历史记录 ---
    if st.session_state.get("analysis_history"):
        st.divider()
        with st.container():
            count = min(5, len(st.session_state.analysis_history))
            with st.expander(f"📚 历史分析记录（最近 {count} 条）"):
                for record in st.session_state.analysis_history[:5]:
                    products_short = record["products"].replace("\n", ", ")[:50]
                    with st.expander(f"{record['time']} — {products_short}"):
                        st.dataframe(pd.DataFrame(record["dataframe"]))
                        st.markdown(clean_markdown(record["report_text"]))


def render_campaign_planning():
    st.markdown('<div class="module-header">'
                '<h1>🎯 活动策划</h1>'
                '<p>输入产品和活动目标，获取完整的活动方案。</p>'
                '</div>', unsafe_allow_html=True)

    col_left, col_right = st.columns(2)

    with col_left:
        product = st.text_input(
            "产品名称",
            placeholder="例：某电商 APP",
        )

    with col_right:
        goal = st.text_input(
            "活动目标",
            placeholder="例：提升新用户注册转化率",
        )

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        temperature = st.slider("创意度", 0.0, 1.5, 0.8, 0.1, key="ctemp")

    with col2:
        max_tokens = st.selectbox("最大长度", [2048, 4096, 8192],
                                  index=1, key="ctokens")

    if st.button("🚀 开始策划", type="primary", use_container_width=True):
        if not product.strip() or not goal.strip():
            st.warning("请填写产品名称和活动目标。")
            return

        with st.spinner("DeepSeek 正在策划中..."):
            placeholder = st.empty()
            collected = ""
            for chunk in chat_stream(
                campaign_planning_system(),
                campaign_planning_user(product.strip(), goal.strip()),
                temperature=temperature,
                max_tokens=max_tokens,
            ):
                collected += chunk
                placeholder.markdown(clean_markdown(collected))


def render_marketing_copy():
    st.markdown('<div class="module-header">'
                '<h1>✍️ 营销文案</h1>'
                '<p>输入产品和目标平台，自动生成适配的营销文案。</p>'
                '</div>', unsafe_allow_html=True)

    col_left, col_right = st.columns(2)

    with col_left:
        product = st.text_input(
            "产品名称",
            placeholder="例：智能手表 Pro",
        )

    with col_right:
        platform = st.selectbox(
            "目标平台",
            options=[
                "小红书",
                "微信公众号",
                "微博",
                "抖音 / 短视频",
                "B站",
                "知乎",
                "官方网站 / 产品详情页",
                "邮件营销 (EDM)",
                "百度竞价 / SEM",
            ],
        )

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        temperature = st.slider("创意度", 0.0, 1.5, 0.8, 0.1, key="mtemp")

    with col2:
        max_tokens = st.selectbox("最大长度", [2048, 4096, 8192],
                                  index=1, key="mtokens")

    if st.button("🚀 生成文案", type="primary", use_container_width=True):
        if not product.strip():
            st.warning("请填写产品名称。")
            return

        with st.spinner("DeepSeek 正在创作中..."):
            placeholder = st.empty()
            collected = ""
            for chunk in chat_stream(
                marketing_copy_system(),
                marketing_copy_user(product.strip(), platform),
                temperature=temperature,
                max_tokens=max_tokens,
            ):
                collected += chunk
                placeholder.markdown(clean_markdown(collected))


def render_weekly_report():
    st.markdown('<div class="module-header">'
                '<h1>📋 周报生成</h1>'
                '<p>输入本周工作内容，自动生成结构化周报。</p>'
                '</div>', unsafe_allow_html=True)

    work_content = st.text_area(
        "描述本周工作内容",
        placeholder="例：\n• 完成用户增长实验 A/B 测试方案设计\n"
                    "• 上线了新版注册流程，转化率提升 12%\n"
                    "• 处理了 3 个线上 bug，已全部修复\n"
                    "• 配合市场部完成双十一活动素材",
        height=200,
    )

    col1, col2, col3 = st.columns([1, 1, 3])
    with col1:
        temperature = st.slider("创意度", 0.0, 1.5, 0.5, 0.1, key="wtemp")

    with col2:
        max_tokens = st.selectbox("最大长度", [2048, 4096, 8192],
                                  index=1, key="wtokens")

    if st.button("🚀 生成周报", type="primary", use_container_width=True):
        if not work_content.strip():
            st.warning("请填写工作内容。")
            return

        with st.spinner("DeepSeek 正在生成中..."):
            placeholder = st.empty()
            collected = ""
            for chunk in chat_stream(
                weekly_report_system(),
                weekly_report_user(work_content.strip()),
                temperature=temperature,
                max_tokens=max_tokens,
            ):
                collected += chunk
                placeholder.markdown(clean_markdown(collected))


def build_context_with_sources() -> str:
    """从 knowledge_docs 构建带来源标记的上下文文本。"""
    docs = st.session_state.get("knowledge_docs", {})
    blocks = []
    for filename, data in docs.items():
        text = data if isinstance(data, str) else data.get("text", "")
        block = f"""
【来源文件】{filename}
【内容开始】
{text}
【内容结束】"""
        blocks.append(block)
    return "\n\n".join(blocks)


def render_knowledge_base():
    st.markdown('<div class="module-header">'
                '<h1>📚 竞品知识库</h1>'
                '<p>上传竞品 PDF 资料，构建知识库并提问。</p>'
                '</div>', unsafe_allow_html=True)

    if "knowledge_docs" not in st.session_state:
        st.session_state["knowledge_docs"] = {}
        if "vector_store" not in st.session_state:
            st.session_state["vector_store"] = VectorStore()

    uploaded_files = st.file_uploader(
        "上传竞品资料（PDF 格式，可多选）",
        accept_multiple_files=True,
        type=["pdf"],
    )

    if uploaded_files:
        for file in uploaded_files:
            if file.name not in st.session_state["knowledge_docs"]:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = "".join(
                        (page.extract_text() or "") + "\n"
                        for page in pdf_reader.pages
                    )
                    st.session_state["knowledge_docs"][file.name] = {
                        "text": text, "source": file.name
                    }
                    chunks = split_text(text)
                    chunk_meta = [{"source": file.name} for _ in chunks]
                    st.session_state["vector_store"].add(chunks, chunk_meta)
                    st.success(f"✅ 已加载: {file.name}")
                except Exception as e:
                    st.error(f"解析失败: {file.name} — {e}")

    if st.session_state["knowledge_docs"]:
        st.subheader("📁 已上传文件")
        for name, data in st.session_state["knowledge_docs"].items():
            text = data if isinstance(data, str) else data.get("text", "")
            char_count = len(text)
            with st.expander(f"📄 {name}（{char_count} 字符）"):
                st.text(text[:2000] + ("..." if char_count > 2000 else ""))
    else:
        st.info("尚未上传文件，请上传 PDF 资料。")

    st.subheader("❓ 提问")
    question = st.text_input(
        "输入你关于竞品的问题（基于已上传的资料）",
        placeholder="例：ChatGPT 和 Claude 最大的差异是什么？",
    )

    col1, _, _ = st.columns([1, 1, 3])
    with col1:
        temperature = st.slider("创意度", 0.0, 1.5, 0.3, 0.1, key="ktemp")

    if st.button("🚀 提问", type="primary", use_container_width=True):
        if not question.strip():
            st.warning("请输入问题。")
            return

        if not st.session_state["knowledge_docs"]:
            st.warning("知识库为空，请先上传 PDF 文件。")
            return

        vs = st.session_state.get("vector_store")
        rag_context = None
        use_rag = False
        try:
            if vs and vs.texts:
                results = vs.search(question.strip(), top_k=3)
                if results:
                    scores = [item.get("score", 0) or 0 for item in results]
                    max_score = max(scores)
                    use_rag = True
                    rag_context = "\n\n".join(
                            f"【来源】{item['meta'].get('source', 'unknown')}\n{item['content']}"
                            for item in results
                        )
        except Exception:
            pass  # 检索失败时降级到直接问答

        with st.spinner("DeepSeek 正在回答中..."):
            placeholder = st.empty()
            collected = ""
            if use_rag and rag_context:
                sys = (
                    "你是专业的 AI 产品分析师。"
                    "请基于以下知识库内容严格回答问题，"
                    "不得使用未提供的信息。\n\n"
                    f"{rag_context}"
                )
                usr = question.strip()
            else:
                sys = "You are a helpful AI assistant. Answer the user directly and concisely in Chinese."
                usr = question.strip()
            for chunk in chat_stream(
                sys,
                usr,
                temperature=temperature,
                max_tokens=4096,
            ):
                collected += chunk
                placeholder.markdown(clean_markdown(collected))

        # Fallback: 如果 RAG 但 LLM 返回空，重试直接问答
        if use_rag and not collected.strip():
            placeholder2 = st.empty()
            fallback = ""
            for chunk in chat_stream(
                "You are a helpful AI assistant. Answer the user directly and concisely in Chinese.",
                question.strip(),
                temperature=temperature,
                max_tokens=4096,
            ):
                fallback += chunk
                placeholder2.markdown(clean_markdown(fallback))


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def main():
    # 检查 API Key
    if not os.getenv("DEEPSEEK_API_KEY"):
        st.warning(
            "⚠️ 未检测到 DeepSeek API Key。\n\n"
            "请在项目根目录创建 `.env` 文件并添加：\n\n"
            "```\nDEEPSEEK_API_KEY=你的密钥\n```\n\n"
            "或者在左侧 Sidebar 的「API 配置」中临时输入。"
        )

    selected = render_sidebar()

    # 路由到对应模块
    if selected == "competitive_analysis":
        render_competitive_analysis()
    elif selected == "campaign_planning":
        render_campaign_planning()
    elif selected == "marketing_copy":
        render_marketing_copy()
    elif selected == "weekly_report":
        render_weekly_report()
    elif selected == "knowledge_base":
        render_knowledge_base()


if __name__ == "__main__":
    main()
